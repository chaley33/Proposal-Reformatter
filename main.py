from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
import os



def remove_column(table, column, index):
    for cell in column.cells:
        cell._tc.getparent().remove(cell._tc)

    grid = table._tbl.find("w:tblGrid", table._tbl.nsmap)
    col_elem = grid[index]
    grid.remove(col_elem)



f = open('path.dat', 'r')
content = f.readline().rstrip()
f.close()

print('Content:' + content)

if content == '':
    path = input("Enter directory path:\n")
else:
    path = content


path = path.replace('\"', '')

document = Document(path)

document.paragraphs[0].add_run().add_picture(os.getcwd() + '/resources/Agfa_logo.png')

for table in document.tables:
    for cell in table._cells:
        for paragraph in cell.paragraphs:
            if paragraph.text.startswith('for Service Plan Coverage'):
                paragraph.text = ''
                run = paragraph.add_run('for Service Plan Coverage')
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x30, 0x3a, 0x46)
            if paragraph.text.startswith('and Conditions'):
                paragraph.text = ''

for table in document.tables:
    for cell in table._cells:
        for paragraph in cell.paragraphs:
            if 'Select Agreement Period:' in paragraph.text:
                table._element.getparent().remove(table._element)

for table in document.tables:
    for cell in table._cells:
        for paragraph in cell.paragraphs:
            if 'Note: Equipment or Software that has reached end of service life will be serviced on a "best effort" basis, but service might no longer be available and replacement equipment or software may need to be purchased by the customer if spare parts are no longer available.' in paragraph.text:
                paragraph.text = 'Company shall continue to provide Service Maintenance Agreement for successive, automatically renewable one (1) year periods (“Renewal Terms”) unless either party provides the other party with written notice of termination of current Services Maintenance Agreement no less than three (3) months prior to the end of the applicable Renewal Term. Notwithstanding the foregoing, Company may suspend Service Maintenance Agreement for nonpayment of any sums owed to Company which are undisputed and ninety (90) days or more past due. The Maintenance Fee for a partial month’s services will be prorated on the basis of a thirty-day (30) month.  Final billing will be subject to all applicable taxes. Each year during the Initial Maintenance Term and each Renewal Term, Company may, increase the Maintenance Fee for any Software or hardware once a year by the greater of (a) four percent (4%) or (b) the annual percentage increase in the CPI Index during the previous twelve (12) month period to include any new assets acquired by the customer. In addition, with respect to Software whose license fee is based on an annual Exam volume, the Maintenance Fee for such Software is subject to an annual increase proportional to the increase in the Customer’s actual Exam volume in the prior twelve (12) month period.  Equipment or Software that has reached end of service life will be serviced on a "best effort" basis, but service might no longer be available and replacement equipment or software may need to be purchased by the customer if spare parts are no longer available.'
    
for table in document.tables:
    for cell in table._cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if run.font.size != None and run.font.size.pt == 4.5:
                    run.font.size = Pt(7.5)

for table in document.tables:
    for index, column in enumerate(table.columns):
        for cell in column.cells:
            for paragraph in cell.paragraphs:
                if 'Discount' in paragraph.text:
                    remove_column(table, column, index)
                        
                

document.save(path)
print("Document successfully edited.")